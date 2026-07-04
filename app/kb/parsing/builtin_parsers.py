"""内置解析器：md / txt / xlsx / docx / pdf(pypdf 兜底)。

不依赖模型下载，作为解析链的可靠保底级；
版面质量要求高的 PDF 应走 MinerU / Docling（见对应 parser）。
"""

import io

from app.kb.parsing.base import Block, ParsedDocument, ParserUnavailable, markdown_to_blocks


class MarkdownParser:
    """Markdown / 纯文本解析器。"""

    name = "builtin_markdown"

    async def parse(self, filename: str, content: bytes) -> ParsedDocument:
        text = content.decode("utf-8", errors="replace")
        if filename.lower().endswith(".txt"):
            # 纯文本：空行分段，无结构
            blocks = [
                Block(type="paragraph", text=p.strip())
                for p in text.replace("\r\n", "\n").split("\n\n")
                if p.strip()
            ]
        else:
            blocks = markdown_to_blocks(text, parser=self.name)
        return ParsedDocument(blocks=blocks, parser=self.name)


class XlsxParser:
    """Excel 解析器：每个 sheet → 标题（sheet 名）+ Markdown 表格块。"""

    name = "builtin_xlsx"

    async def parse(self, filename: str, content: bytes) -> ParsedDocument:
        try:
            from openpyxl import load_workbook
        except ImportError as exc:  # pragma: no cover
            raise ParserUnavailable("openpyxl 未安装") from exc

        wb = load_workbook(io.BytesIO(content), read_only=True, data_only=True)
        blocks: list[Block] = []
        for sheet in wb.worksheets:
            rows = [
                ["" if c is None else str(c).strip() for c in row]
                for row in sheet.iter_rows(values_only=True)
            ]
            rows = [r for r in rows if any(r)]
            if not rows:
                continue
            # sheet 名作为章节标题（切分器会注入标题路径）
            blocks.append(Block(type="heading", text=sheet.title, level=1))
            width = max(len(r) for r in rows)
            md_rows = ["| " + " | ".join(r + [""] * (width - len(r))) + " |" for r in rows]
            md_rows.insert(1, "|" + "---|" * width)
            blocks.append(
                Block(type="table", text="\n".join(md_rows), meta={"rows": len(rows)})
            )
        wb.close()
        return ParsedDocument(blocks=blocks, parser=self.name)


class DocxParser:
    """Word 解析器（python-docx）：保留标题层级与表格。"""

    name = "builtin_docx"

    async def parse(self, filename: str, content: bytes) -> ParsedDocument:
        try:
            import docx
        except ImportError as exc:  # pragma: no cover
            raise ParserUnavailable("python-docx 未安装") from exc

        document = docx.Document(io.BytesIO(content))
        blocks: list[Block] = []
        # 段落（含标题样式识别）
        for para in document.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            style = (para.style.name or "").lower() if para.style else ""
            if style.startswith("heading"):
                try:
                    level = int(style.split()[-1])
                except ValueError:
                    level = 1
                blocks.append(Block(type="heading", text=text, level=min(level, 6)))
            else:
                blocks.append(Block(type="paragraph", text=text))
        # 表格转 Markdown
        for table in document.tables:
            rows = [[c.text.strip() for c in row.cells] for row in table.rows]
            rows = [r for r in rows if any(r)]
            if not rows:
                continue
            width = max(len(r) for r in rows)
            md_rows = ["| " + " | ".join(r + [""] * (width - len(r))) + " |" for r in rows]
            md_rows.insert(1, "|" + "---|" * width)
            blocks.append(Block(type="table", text="\n".join(md_rows), meta={"rows": len(rows)}))
        return ParsedDocument(blocks=blocks, parser=self.name)


class PypdfParser:
    """PDF 纯文本兜底解析器（pypdf）：只保底不保质，无版面/表格还原。"""

    name = "builtin_pypdf"

    async def parse(self, filename: str, content: bytes) -> ParsedDocument:
        try:
            from pypdf import PdfReader
        except ImportError as exc:  # pragma: no cover
            raise ParserUnavailable("pypdf 未安装") from exc

        reader = PdfReader(io.BytesIO(content))
        blocks: list[Block] = []
        for page_no, page in enumerate(reader.pages, start=1):
            text = (page.extract_text() or "").strip()
            for para in text.split("\n\n"):
                para = para.strip()
                if para:
                    blocks.append(Block(type="paragraph", text=para, page=page_no))
        if not blocks:
            raise ParserUnavailable("pypdf 未提取到文本（可能是扫描件，需要 MinerU/Docling OCR）")
        return ParsedDocument(blocks=blocks, parser=self.name)
